"""Word契約書生成モジュール（決定論的コードのみ、AIは使用しない）(v1.4)"""
import os
import re
import shutil
from pathlib import Path

from docx import Document
from dotenv import load_dotenv

load_dotenv()

DEFAULT_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "templates", "contract_template.docx")

# 変数名 → (日本語ラベル, 入力タイプ, 必須) のマッピング
# 入力タイプ: "text" | "textarea" | "date" | "hidden"
VARIABLE_LABEL_MAP: dict[str, tuple[str, str, bool]] = {
    "company_name":    ("顧客名", "text", True),
    "contract_type":   ("契約種別", "hidden", False),   # テンプレートから決定→非表示
    "contract_amount": ("契約金額（例: 1,000,000）", "text", False),
    "start_date":      ("開始日 (YYYY-MM-DD)", "date", False),
    "end_date":        ("終了日 (YYYY-MM-DD)", "date", False),
    "payment_terms":   ("支払条件", "text", False),
    "special_notes":   ("特記事項", "textarea", False),
    "scope":           ("業務範囲", "textarea", False),
    "period":          ("契約期間", "text", False),
    "penalty":         ("違約金条項", "text", False),
    "jurisdiction":    ("管轄裁判所", "text", False),
    "contact_person":  ("担当者名", "text", False),
    "contact_email":   ("担当者メール", "text", False),
    "price":           ("価格", "text", False),
    "quantity":        ("数量", "text", False),
    "delivery_date":   ("納品日 (YYYY-MM-DD)", "date", False),
    "warranty_period": ("保証期間", "text", False),
    "confidentiality": ("秘密保持条項", "textarea", False),
}


def _auto_label(var_name: str) -> str:
    """マッピングにない変数名を自動的に日本語ラベルに変換"""
    return var_name.replace("_", " ").title()


class WordGenerator:
    def __init__(self, template_path: str = None):
        self.template_path = template_path or DEFAULT_TEMPLATE

    def generate(self, params: dict, output_path: str) -> str:
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"テンプレートが見つかりません: {self.template_path}")

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        shutil.copy2(self.template_path, output_path)
        doc = Document(output_path)

        self._replace_in_paragraphs(doc.paragraphs, params)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._replace_in_paragraphs(cell.paragraphs, params)

        doc.save(output_path)
        return os.path.abspath(output_path)

    def _replace_in_paragraphs(self, paragraphs, params: dict):
        for para in paragraphs:
            for key, value in params.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    self._replace_in_runs(para, placeholder, str(value))

    def _replace_in_runs(self, para, placeholder: str, replacement: str):
        full_text = "".join(run.text for run in para.runs)
        if placeholder not in full_text:
            return
        replaced = full_text.replace(placeholder, replacement)
        if para.runs:
            para.runs[0].text = replaced
            for run in para.runs[1:]:
                run.text = ""

    # ------------------------------------------------------------------
    # 変数抽出（v1.4 新規）
    # ------------------------------------------------------------------
    @staticmethod
    def extract_variables(template_path: str) -> list[str]:
        """
        テンプレートから {{変数名}} パターンを全て抽出して返す。
        段落・テーブルセルの両方を走査する。
        """
        if not os.path.exists(template_path):
            return []
        doc = Document(template_path)
        found: set[str] = set()
        pattern = re.compile(r"\{\{(\w+)\}\}")

        for para in doc.paragraphs:
            found.update(pattern.findall(para.text))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found.update(pattern.findall(para.text))

        return sorted(found)

    @staticmethod
    def get_field_info(var_name: str) -> tuple[str, str, bool]:
        """変数名からフォームフィールド情報 (label, type, required) を返す"""
        if var_name in VARIABLE_LABEL_MAP:
            return VARIABLE_LABEL_MAP[var_name]
        return (_auto_label(var_name), "text", False)

    @staticmethod
    def create_sample_template(output_path: str):
        """動作確認用サンプルテンプレートを生成する"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading("業務委託契約書", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph("甲：{{company_name}}（以下「甲」という）")
        doc.add_paragraph("乙：株式会社〇〇（以下「乙」という）")
        doc.add_paragraph()

        doc.add_heading("第1条（契約の目的）", level=1)
        doc.add_paragraph("甲と乙は、以下の条件で{{contract_type}}を締結する。")

        doc.add_heading("第2条（契約期間）", level=1)
        doc.add_paragraph("契約期間は{{start_date}}から{{end_date}}までとする。")

        doc.add_heading("第3条（委託料）", level=1)
        doc.add_paragraph("甲は乙に対し、委託料として{{contract_amount}}円を支払う。")

        doc.add_heading("第4条（支払条件）", level=1)
        doc.add_paragraph("支払条件は{{payment_terms}}とする。")

        doc.add_heading("第5条（特記事項）", level=1)
        doc.add_paragraph("{{special_notes}}")

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
